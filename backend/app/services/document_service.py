"""
ComplyFlow — Document Service

Deterministic file → text and chunk extraction.
NO Gemini calls here — these are pure Python operations.
Gemini is only called in the agent reasoning tools.
"""
from __future__ import annotations

import io
import os
from pathlib import Path
from typing import List, Optional

from pypdf import PdfReader
from docx import Document

from app.services.chunking_service import ChunkedDocument, DocumentChunk, get_chunking_service
from app.services.file_utils import sanitize_filename


class DocumentService:
    """Handles file saving, text extraction, and structured chunking from PDF, DOCX, TXT files."""

    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.chunker = get_chunking_service()

    def save_upload(self, filename: str, content: bytes, project_id: str) -> Path:
        """Save an uploaded file to disk under uploads/{project_id}/."""
        project_dir = Path(self.upload_dir) / project_id
        project_dir.mkdir(parents=True, exist_ok=True)
        safe_name = sanitize_filename(filename or "document")
        file_path = project_dir / safe_name
        file_path.write_bytes(content)
        return file_path



    def extract_chunked_document(
        self,
        filename: str,
        content: bytes,
        document_id: str = "",
    ) -> ChunkedDocument:
        """
        Extract structured, page-aware chunks and metadata from uploaded bytes.
        Detects OCR_REQUIRED for scanned/empty PDFs.
        """
        suffix = Path(filename).suffix.lower()
        doc_id = document_id or Path(filename).stem.replace(" ", "_")

        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return self.chunker.chunk_pdf_pages(
                pages_text=pages,
                document_name=filename,
                document_id=doc_id,
            )

        elif suffix == ".txt":
            raw_text = content.decode("utf-8", errors="replace")
            chunks = self.chunker.chunk_plain_text(
                text=raw_text,
                document_name=filename,
                document_id=doc_id,
                page_number=1,
            )
            return ChunkedDocument(
                document_id=doc_id,
                document_name=filename,
                total_pages=1,
                total_characters=len(raw_text),
                total_chunks=len(chunks),
                status="OK" if raw_text.strip() else "EMPTY",
                diagnostics=f"Extracted {len(chunks)} chunks from TXT.",
                chunks=chunks,
                raw_text=raw_text,
            )

        elif suffix == ".docx":
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            chunks = self.chunker.chunk_plain_text(
                text=raw_text,
                document_name=filename,
                document_id=doc_id,
                page_number=1,
            )
            return ChunkedDocument(
                document_id=doc_id,
                document_name=filename,
                total_pages=1,
                total_characters=len(raw_text),
                total_chunks=len(chunks),
                status="OK" if raw_text.strip() else "EMPTY",
                diagnostics=f"Extracted {len(chunks)} chunks from DOCX.",
                chunks=chunks,
                raw_text=raw_text,
            )

        else:
            raise ValueError(
                f"Unsupported file format: '{suffix}'. "
                f"Supported formats: .pdf, .txt, .docx"
            )

    def extract_text(self, file_path: Path) -> str:
        """
        Extract plain text from a file path.
        Supports: .pdf, .txt, .docx
        """
        content = file_path.read_bytes()
        chunked = self.extract_chunked_document(file_path.name, content)
        if chunked.status == "OCR_REQUIRED":
            return "[OCR_REQUIRED: Scanned PDF contains no extractable text]"
        return chunked.raw_text

    def extract_text_from_bytes(self, filename: str, content: bytes) -> str:
        """Extract text directly from bytes without saving to disk."""
        chunked = self.extract_chunked_document(filename, content)
        if chunked.status == "OCR_REQUIRED":
            raise ValueError(
                f"Document '{filename}' appears to be a scanned or image-only PDF. "
                "OCR_REQUIRED: Text could not be extracted directly."
            )
        if not chunked.raw_text.strip() and chunked.status == "EMPTY":
            raise ValueError(f"Document '{filename}' contains no readable text.")
        return chunked.raw_text

    def validate_file(self, filename: str, content: bytes) -> Optional[str]:
        """
        Validate a file before processing.
        Returns an error message string if invalid, None if valid.
        """
        suffix = Path(filename).suffix.lower()
        supported = {".pdf", ".txt", ".docx"}

        if suffix not in supported:
            return f"Unsupported file format '{suffix}'. Please upload PDF, TXT, or DOCX."

        max_size_mb = 50  # Generous limit for enterprise PDFs
        if len(content) > max_size_mb * 1024 * 1024:
            return f"File '{filename}' exceeds the {max_size_mb}MB size limit."

        if len(content) == 0:
            return f"File '{filename}' is empty."

        return None  # valid
